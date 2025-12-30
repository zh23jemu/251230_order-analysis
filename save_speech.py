from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('F布行出库效率提升分析报告演讲稿', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加演讲信息
info = doc.add_paragraph()
info.add_run('演讲人：XXX').bold = True
info.add_run('\n日期：2025年12月30日')
info.add_run('\n主题：F布行出库效率提升系统性解决方案')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加分隔线
doc.add_page_break()

# 一、开场白
heading1 = doc.add_heading('一、开场白', level=1)
para1 = doc.add_paragraph()
para1.add_run('尊敬的各位领导、同事：').bold = True
para1.add_run('\n\n大家好！今天很荣幸能为大家分享关于F布行出库效率提升的分析报告。在当前激烈的市场竞争环境下，高效的仓储物流管理已成为企业提升竞争力的关键因素之一。出库效率直接影响到客户满意度、运营成本和库存周转率。')
para1.add_run('\n\n本次分析基于F布行2024年全年的订单数据，通过科学的数据分析方法，我们发现了销售的季节性特点、客户下单规律，并运用累托法则和EIQ分析法对SKU和客户进行了分类，最终提出了一套系统性的解决方案。')
para1.add_run('\n\n接下来，我将从数据概况、分析结果、解决方案和预期效果四个方面为大家详细讲解。')

# 二、数据概况
heading2 = doc.add_heading('二、数据概况', level=1)
para2 = doc.add_paragraph()
para2.add_run('我们的分析基于2024年1月至12月的完整订单数据，共包含').bold = True
para2.add_run('571,169条订单').italic = True
para2.add_run('，涉及').bold = True
para2.add_run('101个客户').italic = True
para2.add_run('和').bold = True
para2.add_run('4,000个SKU').italic = True
para2.add_run('。数据涵盖了订单编号、客户编号、SKU编号、订货量和时间等核心字段。')
para2.add_run('\n\n通过对这些数据的深入分析，我们可以全面了解F布行的销售特点和出库需求，为提升出库效率提供数据支撑。')

# 三、数据分析结果
heading3 = doc.add_heading('三、数据分析结果', level=1)

# 3.1 季节性销售特点
heading3_1 = doc.add_heading('1. 季节性销售特点', level=2)
para3_1 = doc.add_paragraph()
para3_1.add_run('月度销售趋势：').bold = True
para3_1.add_run('\n- 销售高峰出现在12月（178,600件）和11月（174,677件）')
para3_1.add_run('\n- 销售低谷出现在8月（93,623件）')
para3_1.add_run('\n- 从1月到6月呈现上升趋势，7-9月有所下降，10-12月再次上升')

para3_2 = doc.add_paragraph()
para3_2.add_run('季度销售分布：').bold = True
para3_2.add_run('\n- 第四季度销售最强（491,958件），占全年销售的28.4%')
para3_2.add_run('\n- 第二季度次之（430,621件），占全年销售的24.7%')
para3_2.add_run('\n- 第一季度和第三季度销售相对较弱')

para3_3 = doc.add_paragraph()
para3_3.add_run('图表展示：').bold = True
para3_3.add_run('（指向屏幕上的月度和季度销售图表）')
para3_3.add_run('\n这张图表清晰地展示了我们的销售季节性波动。了解这一特点对于库存管理和人力调度至关重要。')

# 3.2 客户下单规律
heading3_2 = doc.add_heading('2. 客户下单规律', level=2)
para3_4 = doc.add_paragraph()
para3_4.add_run('订单时间分布：').bold = True
para3_4.add_run('\n- 订单主要集中在15:00-17:00时段')
para3_4.add_run('\n- 16:00达到峰值（157单）')
para3_4.add_run('\n- 上午订单量较少，10:00仅10单')

para3_5 = doc.add_paragraph()
para3_5.add_run('图表展示：').bold = True
para3_5.add_run('（指向屏幕上的小时订单分布图表）')
para3_5.add_run('\n这一规律告诉我们，下午是订单处理的高峰期，需要合理安排人力。')

# 3.3 累托法则分析
heading3_3 = doc.add_heading('3. 累托法则（80/20法则）分析', level=2)
para3_6 = doc.add_paragraph()
para3_6.add_run('SKU分类：').bold = True
para3_6.add_run('\n- A类SKU：800个（20%），贡献48.8%的销售额')
para3_6.add_run('\n- B类SKU：1,201个（30%），贡献24.8%的销售额')
para3_6.add_run('\n- C类SKU：1,999个（50%），贡献26.4%的销售额')

para3_7 = doc.add_paragraph()
para3_7.add_run('客户分类：').bold = True
para3_7.add_run('\n- A类客户：20个（19.8%）')
para3_7.add_run('\n- B类客户：31个（30.7%）')
para3_7.add_run('\n- C类客户：50个（49.5%）')

para3_8 = doc.add_paragraph()
para3_8.add_run('图表展示：').bold = True
para3_8.add_run('（指向屏幕上的SKU和客户分类图表）')
para3_8.add_run('\n根据80/20法则，我们可以将精力集中在20%的核心SKU和客户上，实现资源的最优配置。')

# 3.4 EIQ分析
heading3_4 = doc.add_heading('4. EIQ分析', level=2)
para3_9 = doc.add_paragraph()
para3_9.add_run('- 订单量(I)：订单平均总量为4,653.38件')
para3_9.add_run('\n- 品项数(E)：订单平均包含1,289.30个SKU')
para3_9.add_run('\n- 订货量(Q)：SKU平均订货量为2.51件')

para3_10 = doc.add_paragraph()
para3_10.add_run('图表展示：').bold = True
para3_10.add_run('（指向屏幕上的EIQ分析图表）')
para3_10.add_run('\nEIQ分析帮助我们了解了订单的结构特点，为拣货策略优化提供了依据。')

# 四、SARIMA销售预测
heading4 = doc.add_heading('四、SARIMA销售预测', level=1)
para4 = doc.add_paragraph()
para4.add_run('我们使用SARIMA模型对前3个A类SKU进行了未来52周的销售预测。预测结果显示：')
para4.add_run('\n\n- SKU 6248：呈现稳定增长趋势')
para4.add_run('\n- SKU 6303：销售波动较大')
para4.add_run('\n- SKU 124：销售趋势平稳')
para4.add_run('\n\n这些预测结果可以帮助我们提前规划库存，避免缺货或积压。')

# 五、出库效率提升解决方案
heading5 = doc.add_heading('五、出库效率提升解决方案', level=1)
para5 = doc.add_paragraph()
para5.add_run('基于以上分析，我们提出了以下系统性解决方案：')

# 5.1 基于SKU分类的仓位优化
heading5_1 = doc.add_heading('1. 基于SKU分类的仓位优化', level=2)
para5_1 = doc.add_paragraph()
para5_1.add_run('- A类SKU：放置在离出库口最近的黄金位置')
para5_1.add_run('\n- B类SKU：放置在次优位置')
para5_1.add_run('\n- C类SKU：放置在较远位置')

# 5.2 基于客户分类的拣货策略
heading5_2 = doc.add_heading('2. 基于客户分类的拣货策略', level=2)
para5_2 = doc.add_paragraph()
para5_2.add_run('- A类客户订单优先处理')
para5_2.add_run('\n- 针对不同客户订单特征优化拣货路径')

# 5.3 基于时间分布的人力调度
heading5_3 = doc.add_heading('3. 基于时间分布的人力调度', level=2)
para5_3 = doc.add_paragraph()
para5_3.add_run('- 15:00-17:00高峰期增加拣货人员')
para5_3.add_run('\n- 低谷期安排补货、盘点等工作')

# 5.4 基于预测的库存管理
heading5_4 = doc.add_heading('4. 基于预测的库存管理', level=2)
para5_4 = doc.add_paragraph()
para5_4.add_run('- 根据SARIMA预测结果提前备货')
para5_4.add_run('\n- 优化补货频率和数量')
para5_4.add_run('\n- 对A类SKU设置较高安全库存')

# 5.5 订单分批处理策略
heading5_5 = doc.add_heading('5. 订单分批处理策略', level=2)
para5_5 = doc.add_paragraph()
para5_5.add_run('- 按订单品项数和订货量合理分批')
para5_5.add_run('\n- 对大单和小单采用不同的拣货策略')
para5_5.add_run('\n- 实施波次拣货，提高拣货效率')

# 5.6 仓库布局调整
heading5_6 = doc.add_heading('6. 仓库布局调整', level=2)
para5_6 = doc.add_paragraph()
para5_6.add_run('- 优化拣货路径，减少拣货员行走距离')
para5_6.add_run('\n- 采用分区拣货模式')
para5_6.add_run('\n- 合理规划补货通道和出库通道')

# 六、预期效果
heading6 = doc.add_heading('六、预期效果', level=1)
para6 = doc.add_paragraph()
para6.add_run('实施以上解决方案后，我们预期可以实现：')
para6.add_run('\n\n1. 拣货效率提升30%以上')
para6.add_run('\n2. 库存周转率提高20%')
para6.add_run('\n3. 订单处理时间缩短25%')
para6.add_run('\n4. 客户满意度提升15%')
para6.add_run('\n\n这些效果将直接转化为运营成本的降低和企业竞争力的提升。')

# 七、结论与展望
heading7 = doc.add_heading('七、结论与展望', level=1)
para7 = doc.add_paragraph()
para7.add_run('各位领导、同事，以上就是我们关于F布行出库效率提升的分析报告和解决方案。通过科学的数据分析和系统性的优化建议，我们有信心帮助F布行显著提升出库效率。')
para7.add_run('\n\n当然，这只是一个初步的解决方案，在实施过程中还需要根据实际情况进行调整和优化。我们建议先在小范围内试点，取得经验后再全面推广。')
para7.add_run('\n\n最后，我相信在大家的共同努力下，F布行的仓储物流管理水平一定会迈上一个新的台阶！')
para7.add_run('\n\n谢谢大家！')

# 八、提问环节
heading8 = doc.add_heading('八、提问环节', level=1)
para8 = doc.add_paragraph()
para8.add_run('现在，我愿意回答大家的任何问题。')

# 保存文档
doc.save('F布行出库效率提升分析报告演讲稿.docx')
print('演讲稿已成功保存为Word文档！')